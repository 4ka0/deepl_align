#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys

# Third party packages
import deepl
from environs import Env
from docx import Document


class Segment():
    def __init__(self, source_text, target_text):
        self.source_text = source_text
        self.target_text = target_text


def check_user_input(user_input):

    format_message = (
        "Expected input: python3 translate.py tmx/docx translation.docx glossary.txt\n"
        "(The glossary file is optional.)"
    )

    # Should be 3 or 4 args
    if len(user_input) < 3 or len(user_input) > 4:
        print("Error: Incorrect number of arguments.")
        print(format_message)
        return False, None, None, None

    # 2nd arg should be "tmx" or "docx" specifying the output format
    output_format = user_input[1]
    if output_format != "tmx" and output_format != "docx":
        print('Error: Second argument should be "tmx" or "docx" specifying the output format.')
        print(format_message)
        return False, None, None, None

    # 3rd arg should be a docx file
    translation_file = user_input[2]
    if not translation_file.lower().endswith('.docx'):
        print("Error: Third argument should be the docx file to be translated.")
        print(format_message)
        return False, None, None, None

    # 4th arg, if present, should be a txt file
    if len(user_input) == 4:
        glossary_file = user_input[3]
        if not glossary_file.lower().endswith('.txt'):
            print('Error: Fourth argument should be a ".txt" file.')
            print(format_message)
            return False, None, None, None
    else:
        glossary_file = None

    return True, output_format, translation_file, glossary_file


def setup_deepl_translator():
    env = Env()
    env.read_env()
    auth_key = env.str("AUTH_KEY")
    translator = deepl.Translator(auth_key)
    return translator


def get_source_segments(source_file):
    """
    Reads in text from user-specified docx file.
    File is already split into paragraphs by Document module.
    Text is further split into sentences if a paragraph contains multiple
    sentences.
    """

    print("Extracting text from \"" + source_file + "\".")

    document = Document(source_file)
    segments = []

    for para in document.paragraphs:

        # Split again if paragraph contains multiple sentences.
        if para.text.count("。") >= 2:
            sentences = para.text.split("。")

            for sentence in sentences:
                # Add each single sentence to source_segments list.
                # However, if "。" appears at the end of a string, split()
                # creates an empty string representing the substring that
                # follows "。". Using "if sentence" skips such empty strings.
                # Also, split() removes the "。" delim, so have to add this.
                if sentence:
                    segment = Segment(source_text=sentence + "。", target_text="")
                    segments.append(segment)
        else:
            segment = Segment(source_text=para.text, target_text="")
            segments.append(segment)

    return segments


def get_source_char_count(source_segments):
    source_strings = [segment.source_text for segment in source_segments]
    char_count = sum(len(i) for i in source_strings)
    print("Characters extracted: " + str(char_count))
    return char_count


def check_deepl_usage(source_char_count, translator):
    """
    The monthly limit for the free API is 500000.
    This is set to 499900 to create a buffer of 100 character to allow for
    potential differences in how characters are counted.
    """

    monthly_limit = 499900
    usage = translator.get_usage()

    if source_char_count + usage.character.count >= monthly_limit:
        return False

    return True


def extract_glossary_entries(glossary_file):
    """
    Function to extract entries from user supplied glossary.
    Exits the program if - there is an error when reading the glossary file.
                         - the glossary file does not contain accepted entries.
    """

    print("Reading glossary.")

    entries = {}

    try:
        with open(glossary_file, encoding="utf-8") as f:
            for line in f:

                # Only add to entries if:
                # - there are two entries on a line
                # - both entries are not just whitespace
                # - both entries contain characters
                # Otherwise ignore the line.

                line_entries = line.split("\t")
                if len(line_entries) == 2:
                    if not line_entries[0].isspace() and not line_entries[1].isspace():
                        source = line_entries[0].strip()
                        target = line_entries[1].strip()
                        if source and target:
                            entries[source] = target

    except Exception as e:
        print(
            "An error occurred when reading your glossary file.\n"
            "Error details:"
        )
        print(e)
        sys.exit()

    # Check the length of the entries dict. If zero, output error message and
    # request user to check the formatting of the file content.
    if len(entries) == 0:
        print(
            "Your glossary file did not contain any parsable entries.\n"
            "Please check that the entries in the file are correctly "
            "formatted as a tab-delimited text file."
        )
        sys.exit()

    return entries


def get_filename(whole_file_path):
    """
    Used to get the name of a given file, which is then used to build a name
    for other such as the glossary file created on the DeepL platform and the
    tmx file that is ultimately output.
    """
    filename_plus_extension = os.path.basename(whole_file_path)
    filename = os.path.splitext(filename_plus_extension)[0]
    return filename


def create_deepl_glossary(translator, glossary_name, entries):
    """
    Upload entries to DeepL platform.
    Returns GlossaryInfo object.
    """
    deepl_glossary = translator.create_glossary(
        glossary_name,
        source_lang="JA",
        target_lang="en-US",
        entries=entries,
    )

    return deepl_glossary


def translate_segments(translator, segments, glossary):

    print("Translating text (this may take a little while).")

    # Perform translation using glossary
    if glossary:

        # Get translation for each segment, one segment at a time
        for segment in segments:
            if segment.source_text:
                target_text = translator.translate_text(
                    segment.source_text,
                    source_lang="JA",
                    target_lang="en-US",
                    glossary=glossary,
                )
                segment.target_text = target_text
            else:
                segment.target_text = ""

        # Delete glossary from DeepL platform
        translator.delete_glossary(glossary)

        return segments

    # Perform translation without using glossary
    for segment in segments:
        if segment.source_text:
            target_text = translator.translate_text(
                segment.source_text,
                source_lang="JA",
                target_lang="en-US",
                glossary=None,
            )
            segment.target_text = target_text
        else:
            segment.target_text = ""

    return segments


def create_docx(docx_name, translated_segments):

    docx_filename = docx_name + ".docx"
    docx_path = "output/" + docx_filename

    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    for segment in translated_segments:
        row_cells = table.add_row().cells
        row_cells[0].text = str(segment.source_text)
        row_cells[1].text = str(segment.target_text)

    document.save(docx_path)
    print("DOCX file created in \"" + docx_path + "\".")


def create_tmx(tmx_name, translated_segments):

    tmx_filename = tmx_name + ".tmx"
    tmx_path = "output/" + tmx_filename

    with open(tmx_path, 'w') as f:

        # Write the start of the tmx file
        f.write(
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<!DOCTYPE tmx SYSTEM "tmx11.dtd">\n'
            '<tmx version="1.1">\n'
            '  <header creationtool="DeepL-to-tmx" adminlang="EN-US" datatype="plaintext" '
            'segtype="sentence" srclang="JA"/>\n'
            '  <body>\n'
        )

        # Append the content of each segment to the tmx file
        for segment in translated_segments:
            f.write(
                '    <tu>\n'
                '      <tuv lang="JA">\n'
                '        <seg>' + str(segment.source_text) + '</seg>\n'
                '      </tuv>\n'
                '      <tuv lang="EN-US">\n'
                '        <seg>' + str(segment.target_text) + '</seg>\n'
                '      </tuv>\n'
                '    </tu>\n'
            )

        # Write the end of the tmx file
        f.write(
            '  </body>\n'
            '</tmx>\n\n'
        )

        print("TMX file created in \"" + tmx_path + "\".")


def output_deepl_usage(translator):
    usage = translator.get_usage()
    return (
        "Current DeepL usage for this month: " +
        str(usage.character.count) +
        " (monthly limit: 500000)"
    )


if __name__ == "__main__":

    valid, output_format, source_file, glossary_file = check_user_input(sys.argv)

    if valid:

        translator = setup_deepl_translator()
        source_segments = get_source_segments(source_file)
        source_char_count = get_source_char_count(source_segments)

        if check_deepl_usage(source_char_count, translator):

            if glossary_file:
                glossary_entries = extract_glossary_entries(glossary_file)
                glossary_name = get_filename(glossary_file)
                glossary = create_deepl_glossary(translator, glossary_name, glossary_entries)
                translated_segments = translate_segments(translator, source_segments, glossary)
            else:
                translated_segments = translate_segments(translator, source_segments, None)

        else:
            output_deepl_usage(translator)
            print(
                "The monthly limit has been reached."
                "Please try again next month."
            )
            sys.exit()

        file_name = get_filename(source_file)

        if output_format == "docx":
            create_docx(file_name, translated_segments)
        else:
            create_tmx(file_name, translated_segments)

        print(output_deepl_usage(translator))
